from __future__ import annotations

"""
StudyBuddy Pro — Document Processor

Handles parsing of PDF, DOCX, and TXT files into LangChain Document objects
with rich metadata (source, page, section, chunk_id).
"""

import os
import shutil
from pathlib import Path
from typing import BinaryIO

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config.settings import settings
from utils.logger import get_logger
from utils.helpers import generate_id, format_file_size, sanitize_filename

log = get_logger(__name__)


class DocumentProcessingError(Exception):
    """Raised when a document cannot be processed."""
    pass


class DocumentProcessor:
    """Processes uploaded files into chunked LangChain Documents with metadata."""

    def __init__(self) -> None:
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        settings.ensure_directories()

    def process_file(self, file_path: str | Path,
                     original_name: str | None = None) -> tuple[list[Document], dict]:
        """
        Process a file and return (chunks, file_info).

        Args:
            file_path: Path to the file to process.
            original_name: Original filename (before temp upload renaming).

        Returns:
            Tuple of (list of Document chunks, file metadata dict).

        Raises:
            DocumentProcessingError: If the file cannot be processed.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise DocumentProcessingError(f"File not found: {file_path}")

        name = original_name or file_path.name
        ext = file_path.suffix.lower()
        size = file_path.stat().st_size

        if size == 0:
            raise DocumentProcessingError(f"File '{name}' is empty (0 bytes).")

        if size > settings.max_file_size_mb * 1024 * 1024:
            raise DocumentProcessingError(
                f"File '{name}' exceeds the {settings.max_file_size_mb} MB size limit."
            )

        log.info("Processing file: %s (%s, %s)", name, ext, format_file_size(size))

        try:
            if ext == ".pdf":
                raw_text, page_count, page_texts = self._process_pdf(file_path)
            elif ext == ".docx":
                raw_text, page_count, page_texts = self._process_docx(file_path)
            elif ext == ".txt":
                raw_text, page_count, page_texts = self._process_txt(file_path)
            else:
                raise DocumentProcessingError(
                    f"Unsupported file format: '{ext}'. Supported formats: PDF, DOCX, TXT."
                )
        except DocumentProcessingError:
            raise
        except Exception as e:
            log.exception("Error processing file %s", name)
            raise DocumentProcessingError(
                f"Failed to process '{name}': {str(e)}. The file may be corrupted."
            )

        if not raw_text.strip():
            raise DocumentProcessingError(
                f"No text could be extracted from '{name}'. "
                "The file may be image-based or empty."
            )

        # Chunk with metadata
        chunks = self._create_chunks(raw_text, page_texts, name)

        if not chunks:
            raise DocumentProcessingError(
                f"No chunks could be created from '{name}'. The file may have too little text."
            )

        # Save a copy to uploads dir
        dest = settings.uploads_dir / sanitize_filename(name)
        if not dest.exists():
            shutil.copy2(file_path, dest)

        file_info = {
            "name": name,
            "path": str(dest),
            "pages": page_count,
            "chunks": len(chunks),
            "size": format_file_size(size),
            "size_bytes": size,
            "extension": ext,
        }

        log.info("Processed %s: %d pages, %d chunks", name, page_count, len(chunks))
        return chunks, file_info

    def _process_pdf(self, file_path: Path) -> tuple[str, int, list[dict]]:
        """Extract text from PDF using PyMuPDF."""
        try:
            doc = fitz.open(str(file_path))
        except Exception as e:
            raise DocumentProcessingError(f"Cannot open PDF: {e}")

        page_texts = []
        all_text_parts = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                page_texts.append({
                    "page": page_num + 1,
                    "text": text,
                })
                all_text_parts.append(text)

        page_count = len(doc)
        doc.close()

        if not all_text_parts:
            raise DocumentProcessingError(
                "No readable text found in PDF. It may be scanned/image-based."
            )

        return "\n\n".join(all_text_parts), page_count, page_texts

    def _process_docx(self, file_path: Path) -> tuple[str, int, list[dict]]:
        """Extract text from DOCX using python-docx."""
        try:
            doc = DocxDocument(str(file_path))
        except Exception as e:
            raise DocumentProcessingError(f"Cannot open DOCX: {e}")

        paragraphs = []
        current_section = "Main Content"

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings as sections
            if para.style and para.style.name and "Heading" in para.style.name:
                current_section = text

            paragraphs.append({
                "text": text,
                "section": current_section,
            })

        if not paragraphs:
            # Try tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append({"text": row_text, "section": "Table"})

        full_text = "\n\n".join(p["text"] for p in paragraphs)
        # Approximate page count for DOCX (250 words per page)
        word_count = len(full_text.split())
        page_count = max(1, word_count // 250)

        page_texts = [{"page": 1, "text": full_text}]

        return full_text, page_count, page_texts

    def _process_txt(self, file_path: Path) -> tuple[str, int, list[dict]]:
        """Extract text from plain text file."""
        try:
            text = file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            try:
                text = file_path.read_text(encoding="latin-1")
            except Exception as e:
                raise DocumentProcessingError(f"Cannot read text file: {e}")

        # Approximate page count (50 lines per page)
        lines = text.split("\n")
        page_count = max(1, len(lines) // 50)

        page_texts = [{"page": 1, "text": text}]

        return text, page_count, page_texts

    def _create_chunks(self, full_text: str, page_texts: list[dict],
                       source_name: str) -> list[Document]:
        """Split text into chunks with metadata."""
        documents = []

        # If we have page-level text, chunk per page for better metadata
        if len(page_texts) > 1:
            for page_info in page_texts:
                page_num = page_info["page"]
                page_text = page_info["text"]

                if not page_text.strip():
                    continue

                splits = self.text_splitter.split_text(page_text)

                for i, chunk_text in enumerate(splits):
                    doc = Document(
                        page_content=chunk_text,
                        metadata={
                            "source": source_name,
                            "page": page_num,
                            "chunk_id": f"{source_name}_p{page_num}_c{i}",
                            "chunk_index": i,
                        },
                    )
                    documents.append(doc)
        else:
            # Single-page: split the whole text
            splits = self.text_splitter.split_text(full_text)
            for i, chunk_text in enumerate(splits):
                doc = Document(
                    page_content=chunk_text,
                    metadata={
                        "source": source_name,
                        "page": 1,
                        "chunk_id": f"{source_name}_c{i}",
                        "chunk_index": i,
                    },
                )
                documents.append(doc)

        return documents


# Singleton
document_processor = DocumentProcessor()
